import os
from docx import Document
from docx.shared import Pt

def add_dict_to_doc(document, data, indent=0):
    """
    Recursively adds dictionary data to a docx document.
    Each key is added in bold with an indent based on its level.
    Lists are iterated through; nested dictionaries within lists are processed recursively.
    Null values (None) are handled gracefully.
    """
    indent_space = Pt(indent * 20)
    for key, value in data.items():
        # Create a paragraph for the key
        p_key = document.add_paragraph()
        p_key.paragraph_format.left_indent = indent_space
        run = p_key.add_run(f"{key}:")
        run.bold = True
        
        if isinstance(value, dict):
            # Process nested dictionaries recursively
            add_dict_to_doc(document, value, indent=indent+1)
        elif isinstance(value, list):
            # Process lists by iterating through their items
            for item in value:
                if isinstance(item, dict):
                    add_dict_to_doc(document, item, indent=indent+1)
                else:
                    p_item = document.add_paragraph()
                    p_item.paragraph_format.left_indent = Pt((indent+1) * 20)
                    # If item is None, convert to string "None"
                    p_item.add_run(f"- {item if item is not None else 'None'}")
        else:
            # For simple values, create a new paragraph
            p_value = document.add_paragraph()
            p_value.paragraph_format.left_indent = Pt((indent+1) * 20)
            # Convert None to string "None"
            text = str(value) if value is not None else "Not Found"
            p_value.add_run(text)

def save_multiple_analyses_to_docx(doc_data, folder_name):
    """
    Saves multiple analyses to separate DOCX files based on fund name.
    
    Parameters:
    - doc_data: Dictionary containing the document structure
    - folder_name: Name of the folder to save files in
    
    Returns:
    - List of file paths where documents were saved
    """
    print("doc data", doc_data)
    BASE_DIR = "temp_uploads"
    OUTPUT_DIR = os.path.join(BASE_DIR, folder_name, "docOutputs")
    
    # Create output directory
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
    
    saved_files = []
    
    # Process each analysis item
    if "analysis" in doc_data:  # Changed from doc_data["doc"]
        analyses = doc_data["analysis"]
        
        for analysis in analyses:
            # Extract fund name for file naming
            fund_name = None
            if "GENERAL FUND INFORMATION" in analysis and "Fund Name" in analysis["GENERAL FUND INFORMATION"]:
                fund_name = analysis["GENERAL FUND INFORMATION"]["Fund Name"]
            
            if not fund_name:
                fund_name = f"unknown_fund_{len(saved_files)}"
            
            # Create safe filename
            safe_filename = "".join(c if c.isalnum() or c in "._- " else "_" for c in fund_name)
            file_path = os.path.join(OUTPUT_DIR, f"{safe_filename}.docx")
            
            # Create document
            document = Document()
            document.add_heading(f"Fund Analysis: {fund_name}", level=1)
            
            # Add analysis data
            add_dict_to_doc(document, analysis)
            
            # Save document
            document.save(file_path)
            saved_files.append(file_path)
            print(f"Document saved to {file_path}")
    
    return saved_files

   